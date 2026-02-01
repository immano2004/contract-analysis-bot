"""
Export handler for generating PDF and other format reports
"""
import logging
from typing import Optional
from datetime import datetime
import json

logger = logging.getLogger(__name__)


class ExportHandler:
    """Handle export of analysis results"""
    
    @staticmethod
    def export_to_json(analysis: dict, output_path: Optional[str] = None) -> str:
        """
        Export analysis to JSON format
        
        Args:
            analysis: Analysis dictionary
            output_path: Optional output file path
            
        Returns:
            JSON string or file path
        """
        try:
            json_data = json.dumps(analysis, indent=2, default=str)
            
            if output_path:
                with open(output_path, 'w') as f:
                    f.write(json_data)
                logger.info(f"Analysis exported to {output_path}")
                return output_path
            
            return json_data
        except Exception as e:
            logger.error(f"Error exporting to JSON: {e}")
            raise
    
    @staticmethod
    def export_to_pdf(analysis: dict, output_path: str) -> str:
        """
        Export analysis to PDF format
        
        Args:
            analysis: Analysis dictionary
            output_path: Output file path
            
        Returns:
            File path
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f4788'),
                spaceAfter=30
            )
            elements.append(Paragraph("Contract Analysis Report", title_style))
            elements.append(Spacer(1, 12))
            
            # Metadata
            elements.append(Paragraph(f"<b>Contract Type:</b> {analysis.get('contract_type', 'N/A')}", styles['Normal']))
            elements.append(Paragraph(f"<b>Risk Score:</b> {analysis.get('composite_risk_score', 'N/A')}/100", styles['Normal']))
            elements.append(Paragraph(f"<b>Analysis Date:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Summary
            elements.append(Paragraph("<b>Summary</b>", styles['Heading2']))
            elements.append(Paragraph(analysis.get('summary', 'N/A'), styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Key Risks
            if analysis.get('key_risks'):
                elements.append(Paragraph("<b>Key Risks</b>", styles['Heading2']))
                for risk in analysis['key_risks'][:5]:
                    elements.append(Paragraph(f"• {risk}", styles['Normal']))
                elements.append(Spacer(1, 12))
            
            # Recommendations
            if analysis.get('recommendations'):
                elements.append(Paragraph("<b>Recommendations</b>", styles['Heading2']))
                for rec in analysis['recommendations'][:5]:
                    elements.append(Paragraph(f"✓ {rec}", styles['Normal']))
            
            doc.build(elements)
            logger.info(f"Analysis exported to PDF: {output_path}")
            return output_path
        
        except ImportError:
            logger.warning("reportlab not installed. Install with: pip install reportlab")
            raise
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    @staticmethod
    def export_to_docx(analysis: dict, output_path: str) -> str:
        """
        Export analysis to DOCX format
        
        Args:
            analysis: Analysis dictionary
            output_path: Output file path
            
        Returns:
            File path
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Contract Analysis Report', 0)
            title.runs[0].font.color.rgb = RGBColor(31, 71, 136)
            
            # Metadata
            doc.add_paragraph(f"Contract Type: {analysis.get('contract_type', 'N/A')}")
            doc.add_paragraph(f"Risk Score: {analysis.get('composite_risk_score', 'N/A')}/100")
            doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Summary
            doc.add_heading('Summary', level=2)
            doc.add_paragraph(analysis.get('summary', 'N/A'))
            
            # Key Risks
            if analysis.get('key_risks'):
                doc.add_heading('Key Risks', level=2)
                for risk in analysis['key_risks'][:5]:
                    doc.add_paragraph(risk, style='List Bullet')
            
            # Recommendations
            if analysis.get('recommendations'):
                doc.add_heading('Recommendations', level=2)
                for rec in analysis['recommendations'][:5]:
                    doc.add_paragraph(rec, style='List Number')
            
            doc.save(output_path)
            logger.info(f"Analysis exported to DOCX: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")
            raise
